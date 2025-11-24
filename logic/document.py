from docx import Document
from htmldocx import HtmlToDocx
import io
from docx.shared import Pt, RGBColor

def create_word_from_html(html_content, raw_data=""):
    """
    ქმნის Word-ს: 
    1. HTML (ლამაზი ნაწილი)
    2. Raw Data (დანართი ბოლოში)
    """
    doc = Document()
    new_parser = HtmlToDocx()
    
    # --- ნაწილი 1: ლამაზი რეპორტი ---
    # სათაური
    title = doc.add_heading('სატენდერო ტექნიკური წინადადება', 0)
    title.alignment = 1 # ცენტრში
    
    # HTML-ის ჩაშენება
    new_parser.add_html_to_document(html_content, doc)
    
    # --- ნაწილი 2: დანართი (ნედლი მასალა) ---
    if raw_data:
        doc.add_page_break() # გადავდივართ ახალ გვერდზე
        
        heading = doc.add_heading('დანართი: დამუშავებული წყაროები', 1)
        heading.runs[0].font.color.rgb = RGBColor(100, 100, 100) # ნაცრისფერი სათაური
        
        doc.add_paragraph("ქვემოთ მოცემულია სისტემის მიერ PDF და Excel ფაილებიდან ამოღებული ინფორმაცია, რასაც დაეყრდნო ანალიზი.")
        doc.add_paragraph("_" * 50)
        
        # ვამატებთ ნედლ ტექსტს (ცოტა პატარა შრიფტით, რომ არ ყვიროდეს)
        p = doc.add_paragraph(raw_data)
        p.style.font.size = Pt(9) 

    bio = io.BytesIO()
    doc.save(bio)
    return bio