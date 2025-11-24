from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import pdfplumber
import pandas as pd
from docx import Document
import io
import re
import zipfile
from openai import OpenAI

# --- 1. კონფიგურაცია და უსაფრთხოება ---
st.set_page_config(page_title="Tender AI Pro", page_icon="📂")

if "OPENAI_API_KEY" in st.secrets:
    API_KEY = st.secrets["OPENAI_API_KEY"]
else:
    API_KEY = ""

def check_password():
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False
    if not st.session_state.password_correct:
        st.title("🔒 შესვლა სისტემაში")
        pwd = st.text_input("პაროლი", type="password")
        if st.button("შესვლა"):
            if pwd == st.secrets["APP_PASSWORD"]:
                st.session_state.password_correct = True
                st.rerun()
            else:
                st.error("პაროლი არასწორია")
        return False
    return True

if not check_password():
    st.stop()

# --- 2. დამხმარე ფუნქციები ---

def extract_contact_info(text):
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    phones = re.findall(r'\b5\d{2}[-\s]?\d{2}[-\s]?\d{2}[-\s]?\d{2}\b', text)
    return set(emails), set(phones)

# --- განახლებული ფუნქცია: ლამაზი Word-ის შექმნა ---
def create_word_docx(text_content):
    doc = Document()
    
def create_word_docx(text_content):
    doc = Document()
    
    # --- 1. ძირითადი სტილის შეცვლა (შრიფტი Arial) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- 2. მთავარი სათაური (დიდი და ლურჯი) ---
    title = doc.add_heading('Tender AI - ანალიტიკური რეპორტი', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # სათაურის ფერის შეცვლა (მუქი ლურჯი)
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 51, 102) 
    title_run.font.bold = True

    # ხაზი სათაურის ქვემოთ
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("") # ცარიელი ადგილი

    # ტექსტის დაშლა
    lines = text_content.split('\n')

    for line in lines:
        line = line.strip()
        if not line: continue

        # --- სათაურები (###) ---
        if line.startswith('### ') or line.startswith('## '):
            clean_text = line.replace('#', '').strip()
            heading = doc.add_heading(clean_text, level=2)
            
            # სათაურის სტილი (ლურჯი)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 102, 204) # ღია ლურჯი
            run.font.size = Pt(14)
            run.font.name = 'Arial'
        
        # --- ბულეტები (- ან *) ---
        elif line.startswith('- ') or line.startswith('* '):
            clean_text = line.replace('- ', '').replace('* ', '').strip()
            p = doc.add_paragraph(clean_text, style='List Bullet')
            
            # ბულეტებში გამუქების დამუშავება
            if "**" in clean_text:
                p.clear() # ვშლით და თავიდან ვაწყობთ
                parts = clean_text.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1: # გამუქება
                        run.bold = True
                        run.font.color.rgb = RGBColor(50, 50, 50) # მუქი ნაცრისფერი

        # --- ჩვეულებრივი ტექსტი ---
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6) # დაშორება აბზაცებს შორის

            # გამუქების (**text**) დამუშავება
            if "**" in line:
                parts = line.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    run.font.name = 'Arial'
                    if i % 2 == 1: # გამუქება
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0) # შავი
            else:
                run = p.add_run(line)
                run.font.name = 'Arial'

    # --- ფუტერი (ბოლოში მიაწეროს) ---
    doc.add_paragraph("")
    doc.add_paragraph("_" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer = doc.add_paragraph("დოკუმენტი გენერირებულია Tender AI-ს მიერ")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    bio = io.BytesIO()
    doc.save(bio)
    return bio

    # ფაილის შენახვა მეხსიერებაში
    bio = io.BytesIO()
    doc.save(bio)
    return bio

def ask_ai(full_text):
    if not API_KEY: return "⚠️ API Key არ არის!"
    client = OpenAI(api_key=API_KEY)
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "შენ ხარ ტენდერების ექსპერტი. გააანალიზე ყველა მიწოდებული ფაილი (PDF, Word, Excel) ერთად."},
            {"role": "user", "content": f"აი მასალები:\n\n{full_text[:25000]}"} 
        ]
    )
    return response.choices[0].message.content

# --- 3. ფაილების წამკითხავი ფუნქცია (უნივერსალური) ---
def process_file(file_bytes, file_name):
    """ეს ფუნქცია იღებს ფაილს და აბრუნებს ტექსტს, ტიპის მიხედვით"""
    text_content = ""
    
    try:
        # --> PDF
        if file_name.endswith(".pdf"):
            with pdfplumber.open(file_bytes) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt: text_content += txt + "\n"
        
        # --> WORD (.docx)
        elif file_name.endswith(".docx"):
            doc = Document(file_bytes)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        
        # --> EXCEL
        elif file_name.endswith(".xlsx") or file.name.endswith(".xls"):
            df = pd.read_excel(file_bytes)
            text_content = df.to_string(index=False)
            
    except Exception as e:
        return f"\n[შეცდომა {file_name}-ის კითხვისას: {e}]\n"

    return f"\n\n--- ფაილი: {file_name} ---\n{text_content}"

# --- 4. მთავარი ინტერფეისი ---
st.title("📂 Tender AI - ყველა ფორმატი")
st.write("ატვირთეთ: PDF, Word, Excel ან ZIP არქივი.")

uploaded_files = st.file_uploader(
    "ფაილები", 
    type=["pdf", "xlsx", "xls", "docx", "zip"], 
    accept_multiple_files=True
)

if uploaded_files:
    st.success(f"✅ მიღებულია {len(uploaded_files)} ფაილი")
    combined_text = ""
    
    for file in uploaded_files:
        # თუ ZIP ფაილია - ვხსნით და შიგნით ვიხედებით
        if file.name.endswith(".zip"):
            with zipfile.ZipFile(file) as z:
                for sub_file_name in z.namelist():
                    # ვფილტრავთ სისტემურ ფაილებს (Mac-ის __MACOSX და ა.შ.)
                    if not sub_file_name.startswith("__") and not sub_file_name.endswith("/"):
                        with z.open(sub_file_name) as f:
                            # ფაილს ვკითხულობთ ბაიტებად
                            file_bytes = io.BytesIO(f.read())
                            # ვაგზავნით დასამუშავებლად
                            combined_text += process_file(file_bytes, sub_file_name)
        
        # თუ ჩვეულებრივი ფაილია
        else:
            combined_text += process_file(file, file.name)

    # შედეგების გამოტანა
    emails, phones = extract_contact_info(combined_text)
    with st.sidebar:
        st.header("🔍 ნაპოვნია")
        if emails: st.write("📧", ", ".join(emails))
        if phones: st.write("📱", ", ".join(phones))

    with st.expander("ნახე სრული ტექსტი"):
        st.text(combined_text)
    
    if st.button("✨ გააანალიზე ყველაფერი (AI)"):
        if not API_KEY:
            st.error("API Key არ არის!")
        else:
            with st.spinner("AI აანალიზებს ZIP-ს, Word-ს, PDF-ს და Excel-ს..."):
                try:
                    res = ask_ai(combined_text)
                    st.markdown("### 🤖 ანალიზი:")
                    st.write(res)
                    docx = create_word_docx(res)
                    st.download_button("📥 გადმოწერა", docx, "analysis.docx")
                except Exception as e:
                    st.error(f"შეცდომა: {e}")